#!/usr/bin/env python3
"""
Десктопное приложение для анализа жирнокислотного состава молока (PyQt6)
"""

import os
import sys

# Принудительно устанавливаем платформу Qt (по ОС)
if sys.platform.startswith('linux'):
    os.environ['QT_QPA_PLATFORM'] = 'xcb'
elif sys.platform == 'darwin':
    os.environ['QT_QPA_PLATFORM'] = 'cocoa'
elif sys.platform.startswith('win'):
    os.environ['QT_QPA_PLATFORM'] = 'windows'

from PyQt6.QtWidgets import *
from PyQt6.QtCore import *
from PyQt6.QtGui import *
import numpy as np
import matplotlib

matplotlib.use('QtAgg')
from matplotlib.backends.backend_qtagg import FigureCanvasQTAgg
from matplotlib.figure import Figure

from database import DatabaseManager
from utils import validate_diet_ratios, check_fatty_acid_ranges
from utils.constants import FATTY_ACID_NAMES, ingredient_names, nutrient_names
from preprocessing import (
    parse_pdf_diet,
    get_nutrients_data,
    INGREDIENT_FEATURES,
    NUTRIENT_FEATURES,
)
from preprocessing.parser import numeric_from_str

from ingredient_model import predict_from_ingredients
from nutrient_model import load_model, run_predictions


class MplCanvas(FigureCanvasQTAgg):
    """Холст matplotlib для встраивания графиков"""

    def __init__(self, parent=None, width=8, height=6, dpi=100):
        fig = Figure(figsize=(width, height), dpi=dpi)
        self.axes = fig.add_subplot(111)
        super().__init__(fig)


class ClickableLabel(QLabel):
    """QLabel с сигналом клика по изображению"""

    clicked = pyqtSignal(str)

    def __init__(self, image_path: str, parent=None):
        super().__init__(parent)
        self.image_path = image_path
        self.setCursor(Qt.CursorShape.PointingHandCursor)

    def mousePressEvent(self, event):
        if event.button() == Qt.MouseButton.LeftButton:
            self.clicked.emit(self.image_path)
        super().mousePressEvent(event)


class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.db = DatabaseManager()
        self.distribution_points = None
        self.current_diet_data = {}
        self.current_analysis_data = {}
        self.now_open_file = ""
        self.ing_df_glob = 0
        self.setWindowTitle("Анализ жирнокислотного состава молока")
        self.setGeometry(100, 100, 1400, 900)

        # Хранилище для предсказаний и ID рациона
        self.current_predictions = {}
        self.current_diet_id = None

        # Путь к папке с заготовленными графиками
        base_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), os.pardir))
        self.visuals_dir = os.path.join(base_dir, 'visuals')

        self.init_ui()

    def init_ui(self):
        """Инициализация интерфейса"""
        # Главный виджет
        main_widget = QWidget()
        self.setCentralWidget(main_widget)
        self.nutrients_model = load_model()

        # Главный layout
        main_layout = QVBoxLayout(main_widget)

        # Заголовок
        title = QLabel("Анализ жирнокислотного состава молока")
        title_font = QFont()
        title_font.setPointSize(18)
        title_font.setBold(True)
        title.setFont(title_font)
        title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        main_layout.addWidget(title)

        # Вкладки
        self.tabs = QTabWidget()
        main_layout.addWidget(self.tabs)

        # Создание всех вкладок
        self.create_loading_tab()  # Переименовано из analysis_tab
        self.create_predictions_tab()
        self.create_results_tab()  # Переименовано из data_management_tab
        self.create_about_tab()

        # Статус бар
        self.statusBar().showMessage("Готов к работе")

        # Автозагрузка изображений при переключении на вкладку результатов
        self.tabs.currentChanged.connect(self.on_tab_changed)

    def create_loading_tab(self):
        """Вкладка загрузки данных"""
        tab = QWidget()
        self.tabs.addTab(tab, "Загрузка")

        layout = QHBoxLayout(tab)

        # Левая панель - ввод
        left_panel = QWidget()
        left_layout = QVBoxLayout(left_panel)

        # Загрузка файлов
        file_group = QGroupBox("Загрузка файлов")
        file_layout = QVBoxLayout()

        pdf_btn = QPushButton("Загрузить PDF файл")
        pdf_btn.clicked.connect(self.load_pdf)
        file_layout.addWidget(pdf_btn)

        file_group.setLayout(file_layout)
        left_layout.addWidget(file_group)

        # Ручной ввод
        manual_group = QGroupBox("Ручной ввод данных")
        manual_layout = QVBoxLayout()

        # Название рациона
        name_layout = QHBoxLayout()
        name_layout.addWidget(QLabel("Название рациона:"))
        self.diet_name = QLineEdit("Пользовательский рацион")
        name_layout.addWidget(self.diet_name)
        manual_layout.addLayout(name_layout)

        # Ингредиенты рациона (% СВ - сухого вещества)
        ingr_group = QGroupBox("Ингредиенты рациона (% СВ)")
        ingr_scroll = QScrollArea()
        ingr_scroll.setWidgetResizable(True)
        ingr_widget = QWidget()
        ingr_layout = QGridLayout(ingr_widget)

        self.ingredient_inputs = {}
        # Используем читаемые имена из utils.constants.ingredient_names
        ingredient_items = [(code, ingredient_names.get(code, label), 0.0) for code, label in INGREDIENT_FEATURES]

        for i, (key, label, default) in enumerate(ingredient_items):
            row = i // 2
            col = (i % 2) * 2
            ingr_layout.addWidget(QLabel(f"{label} (% СВ):"), row, col)
            spin = QDoubleSpinBox()
            spin.setRange(0, 100)
            spin.setValue(default)
            spin.setSuffix(" %")
            spin.setMinimumWidth(120)
            spin.setMaximumWidth(150)
            self.ingredient_inputs[key] = spin
            ingr_layout.addWidget(spin, row, col + 1)

        ingr_scroll.setWidget(ingr_widget)
        ingr_scroll.setMinimumHeight(300)
        ingr_scroll.setMaximumHeight(400)

        ingr_group_layout = QVBoxLayout()
        ingr_group_layout.addWidget(ingr_scroll)
        ingr_group.setLayout(ingr_group_layout)
        manual_layout.addWidget(ingr_group)

        # Нутриенты (Value_i признаки)
        nutr_group = QGroupBox("Нутриенты (СВ)")
        nutr_layout = QGridLayout()

        self.nutrient_inputs = {}
        # Показываем читаемые лейблы нутриентов
        nutrient_items = [(feat, nutrient_names.get(feat, feat), 0.0) for feat in NUTRIENT_FEATURES]

        for i, (key, label, default) in enumerate(nutrient_items):
            nutr_layout.addWidget(QLabel(label + ":"), i, 0)
            spin = QDoubleSpinBox()
            spin.setRange(0, 10000)
            spin.setValue(default)
            self.nutrient_inputs[key] = spin
            nutr_layout.addWidget(spin, i, 1)

        nutr_group.setLayout(nutr_layout)
        manual_layout.addWidget(nutr_group)

        # Кнопка сохранения
        save_btn = QPushButton("Сохранить данные")
        save_btn.clicked.connect(self.save_loading_data)
        manual_layout.addWidget(save_btn)

        manual_group.setLayout(manual_layout)
        left_layout.addWidget(manual_group)

        left_layout.addStretch()
        layout.addWidget(left_panel, 1)

        # Правая панель - результаты парсинга
        right_panel = QWidget()
        right_layout = QVBoxLayout(right_panel)

        # Таблица результатов загрузки
        results_group = QGroupBox("Результаты загрузки")
        results_layout = QVBoxLayout()

        self.loading_table = QTableWidget()
        self.loading_table.setColumnCount(3)
        self.loading_table.setHorizontalHeaderLabels([
            'Параметр', 'Значение', 'Статус'
        ])
        self.loading_table.horizontalHeader().setStretchLastSection(True)
        results_layout.addWidget(self.loading_table)

        results_group.setLayout(results_layout)
        right_layout.addWidget(results_group)

        layout.addWidget(right_panel, 2)

    def create_predictions_tab(self):
        """Вкладка предсказаний"""
        tab = QWidget()
        self.tabs.addTab(tab, "Предсказания")

        layout = QVBoxLayout(tab)

        # Кнопка
        pred_btn = QPushButton("Сгенерировать предсказания")
        pred_btn.clicked.connect(self.generate_predictions)
        layout.addWidget(pred_btn)

        # Таблица предсказаний
        pred_group = QGroupBox("Результаты предсказаний")
        pred_layout = QVBoxLayout()

        self.pred_table = QTableWidget()
        self.pred_table.setColumnCount(4)
        self.pred_table.setHorizontalHeaderLabels([
            'Жирная кислота', 'Предсказанное значение (%)', 'Уровень по ГОСТу', 'Уверенность'
        ])
        self.pred_table.horizontalHeader().setStretchLastSection(True)
        pred_layout.addWidget(self.pred_table)

        pred_group.setLayout(pred_layout)
        layout.addWidget(pred_group)

    def create_results_tab(self):
        """Вкладка управления результатами"""
        tab = QWidget()
        self.results_tab_index = self.tabs.addTab(tab, "Управление результатами")

        layout = QVBoxLayout(tab)

        # Кнопки
        btn_layout = QHBoxLayout()

        export_docx_btn = QPushButton("Экспорт в DOCX")
        export_docx_btn.clicked.connect(self.export_to_docx)
        btn_layout.addWidget(export_docx_btn)

        export_pdf_btn = QPushButton("Экспорт в PDF")
        export_pdf_btn.clicked.connect(self.export_to_pdf)
        btn_layout.addWidget(export_pdf_btn)

        print_btn = QPushButton("Печать")
        print_btn.clicked.connect(self.print_results)
        btn_layout.addWidget(print_btn)

        layout.addLayout(btn_layout)

        # Таблица результатов предсказаний
        results_group = QGroupBox("Построить графики")
        results_layout = QVBoxLayout()

        self.results_pred_table = QTableWidget()
        self.results_pred_table.setColumnCount(4)
        self.results_pred_table.setHorizontalHeaderLabels([
            'Жирная кислота', 'Значение (%)', 'Уровень по ГОСТу', 'Уверенность'
        ])

        # Устанавливаем, чтобы столбцы были равной ширины и растягивались
        header = self.results_pred_table.horizontalHeader()
        header.setSectionResizeMode(QHeaderView.ResizeMode.Stretch)

        results_group.setLayout(results_layout)
        layout.addWidget(results_group)

        # Готовые графики
        graph_group = QGroupBox("")
        graph_layout = QVBoxLayout()

        # Кнопки для графиков
        graph_btn_layout = QHBoxLayout()
        dist_btn = QPushButton("Показать гистограмму")
        dist_btn.clicked.connect(self.show_hist)
        graph_btn_layout.addWidget(dist_btn)

        trend_btn = QPushButton("Показать тренды")
        trend_btn.clicked.connect(self.show_trends)
        graph_btn_layout.addWidget(trend_btn)

        graph_layout.addLayout(graph_btn_layout)

        self.canvas = MplCanvas(self, width=10, height=6)
        graph_layout.addWidget(self.canvas)

        graph_group.setLayout(graph_layout)
        layout.addWidget(graph_group)

        # Галерея
        prepared_group = QGroupBox("Графики")
        prepared_layout = QVBoxLayout()

        prepared_btns = QHBoxLayout()
        refresh_btn = QPushButton("Обновить список")
        refresh_btn.clicked.connect(self.load_prepared_images)
        prepared_btns.addWidget(refresh_btn)
        prepared_btns.addStretch()
        prepared_layout.addLayout(prepared_btns)

        self.prepared_scroll = QScrollArea()
        self.prepared_scroll.setWidgetResizable(True)
        self.prepared_container = QWidget()
        self.prepared_grid = QGridLayout(self.prepared_container)
        self.prepared_grid.setContentsMargins(6, 6, 6, 6)
        self.prepared_grid.setHorizontalSpacing(12)
        self.prepared_grid.setVerticalSpacing(12)
        self.prepared_scroll.setWidget(self.prepared_container)
        prepared_layout.addWidget(self.prepared_scroll)

        prepared_group.setLayout(prepared_layout)
        layout.addWidget(prepared_group)

        # Первая загрузка миниатюр
        self.load_prepared_images()

    def create_about_tab(self):
        """Вкладка о программе"""
        tab = QWidget()
        self.tabs.addTab(tab, "О программе")

        layout = QVBoxLayout(tab)

        # Информация
        info_text = QTextEdit()
        info_text.setReadOnly(True)
        info_text.setHtml("""
        <h2>Инструмент анализа жирнокислотного состава молока</h2>
        <p>Это приложение помогает анализировать и предсказывать жирнокислотный состав молока на основе параметров рациона.</p>
        <h3>Возможности:</h3>
        <ul>
            <li><b>Загрузка</b>: Загрузка Excel/PDF файлов или ручной ввод данных</li>
            <li><b>Предсказания</b>: Генерация предсказаний с использованием ML моделей</li>
            <li><b>Управление результатами</b>: Визуализация данных с интерактивными графиками, экспорт и печать</li>
        </ul>
        <h3>Технологический стек:</h3>
        <ul>
            <li><b>GUI</b>: PyQt6</li>
            <li><b>База данных</b>: SQLite</li>
            <li><b>Визуализация</b>: Matplotlib</li>
            <li><b>Обработка данных</b>: Pandas, NumPy</li>
        </ul>
        """)
        layout.addWidget(info_text)

    # Обработчики событий
    def load_pdf(self):
        """Загрузка PDF файла"""
        file_path, _ = QFileDialog.getOpenFileName(
            self, "Выберите PDF файл", "", "PDF files (*.pdf)"
        )
        self.now_open_file = file_path
        if file_path:
            ing_df, nut_df = parse_pdf_diet(file_path)
            self.ing_df_glob = ing_df.copy()
            # Также в форму ингредиентов по кодам
            # Отображаем результаты парсинга
            # Подставляем распознанные значения в поля ввода
            try:
                self._populate_inputs_from_loaded(ing_df, nut_df)
            except Exception:
                pass
            self.display_loading_results(ing_df, nut_df, "PDF")

    def display_loading_results(self, df_ingr, df_nutr, file_type):
        all_items = []

        def to_float(x):
            try:
                return float(x)
            except Exception:
                return numeric_from_str(x)

        if df_ingr is not None and not df_ingr.empty:
            for col in df_ingr.columns:
                key = str(col)
                series = df_ingr[col].dropna()
                if len(series) == 0:
                    value = None
                elif len(series) == 1:
                    value = series.iloc[0]
                else:
                    try:
                        value = float(series.mean())
                    except Exception:
                        value = series.iloc[0]
                v = to_float(value)
                if v is not None and v != 0:
                    all_items.append((key, v))

        if df_nutr is not None and not df_nutr.empty:
            for col in [c for c in df_nutr.columns if c in nutrient_names]:
                key = str(col)
                series = df_nutr[col].dropna()
                if len(series) == 0:
                    value = None
                elif len(series) == 1:
                    value = series.iloc[0]
                else:
                    try:
                        value = float(series.mean())
                    except Exception:
                        value = series.iloc[0]
                v = to_float(value)
                if v is not None and v != 0:
                    all_items.append((key, v))

        self.loading_table.setRowCount(len(all_items))
        self.loading_table.setColumnCount(3)
        self.loading_table.setHorizontalHeaderLabels(["Параметр", "Значение", "Статус"])
        for row, (key, v) in enumerate(all_items):
            # Показываем русские названия нутриентов, если доступны
            if key in nutrient_names:
                param_name = nutrient_names[key]
            else:
                param_name = FATTY_ACID_NAMES.get(key, key) if 'FATTY_ACID_NAMES' in globals() else key
            self.loading_table.setItem(row, 0, QTableWidgetItem(param_name))
            self.loading_table.setItem(row, 1, QTableWidgetItem(f"{float(v):.2f}"))
            self.loading_table.setItem(row, 2, QTableWidgetItem("OK"))
        self.loading_table.resizeColumnsToContents()

    def _populate_inputs_from_loaded(self, df_ingr, df_nutr):
        """Заполнить спинбоксы ингредиентов и нутриентов данными после загрузки.

        df_ingr: DataFrame, где колонки — читаемые названия из feed_types (INGREDIENT_FEATURES)
        df_nutr: DataFrame, где колонки — Value_i
        """
        # Ингредиенты: маппим label -> code и подставляем проценты СВ
        try:
            if df_ingr is not None and not df_ingr.empty:
                label_to_code = {label: code for code, label in INGREDIENT_FEATURES}
                row0 = df_ingr.iloc[0]
                for label, value in row0.items():
                    code = label_to_code.get(label)
                    if not code:
                        continue
                    spin = self.ingredient_inputs.get(code)
                    if spin is None:
                        continue
                    try:
                        val = float(value)
                        if not np.isnan(val):
                            spin.setValue(val)
                    except Exception:
                        pass
        except Exception:
            pass

        # Нутриенты: подставляем только те Value_i, которые отображаются в форме
        try:
            if df_nutr is not None and not df_nutr.empty:
                row0 = df_nutr.iloc[0]
                for key, spin in self.nutrient_inputs.items():
                    if key not in row0.index:
                        continue
                    raw_val = row0[key]
                    val = None
                    try:
                        # сначала пробуем напрямую
                        val = float(raw_val)
                    except Exception:
                        # затем извлекаем число из строки вида '1,23 %'
                        try:
                            val = numeric_from_str(raw_val)
                        except Exception:
                            val = None
                    if val is not None and not np.isnan(val):
                        spin.setValue(val)
        except Exception:
            pass

    def save_loading_data(self):
        """Сохранение данных из вкладки загрузки"""
        try:
            # Сбор данных из ингредиентов и нутриентов
            self.current_diet_data = {k: v.value() for k, v in self.ingredient_inputs.items()}
            nutrients_data = {k: v.value() for k, v in self.nutrient_inputs.items()}
            # Валидация рациона
            is_valid, message = validate_diet_ratios(self.current_diet_data)
            if not is_valid:
                QMessageBox.warning(self, "Ошибка валидации", message)
                return
            entered_codes = {code: val for code, val in self.current_diet_data.items() if val > 0}
            diet_id = self.db.add_diet(
                name=self.diet_name.text(),
            )
            # Сохраняем diet_id для использования в предсказаниях
            self.current_diet_id = diet_id
            QMessageBox.information(self, "Успех",
                                    "Данные сохранены в БД! Теперь можете перейти в вкладку Предсказания.")
            self.statusBar().showMessage(f"Данные сохранены (ID рациона: {diet_id})")
            # Переключаемся на вкладку предсказаний
            self.tabs.setCurrentIndex(1)
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка сохранения: {str(e)}")

    def generate_predictions(self):
        if not self.now_open_file:
            QMessageBox.warning(
                self, "Нет данных",
                "Сначала загрузите файл с рационом во вкладке «Загрузка»."
            )
            return
        try:
            # 2) По нутриентам: соберём Value_i
            nutrients_data_no_vibecode = get_nutrients_data(self.now_open_file)
            print(nutrients_data_no_vibecode, len(nutrients_data_no_vibecode))
            nutrients_by_feat = {k: v.value() for k, v in self.nutrient_inputs.items() if v.value() > 0}
            # Предсказания двух моделей
            pred_ingr = predict_from_ingredients(self.ing_df_glob).tolist()[0]
            pred_nutr = run_predictions(nutrients_data_no_vibecode, self.nutrients_model).tolist()[0]
            print(pred_ingr, pred_nutr)
            # Усреднение
            acids = ['Масляная', 'Капроновая', 'Каприловая', 'Каприновая', 'Деценовая', 'Лауриновая',
                     'Миристиновая', 'Миристолеиновая', 'Пальмитиновая', 'Пальмитолеиновая',
                     'Стеариновая', 'Олеиновая', 'Линолевая', 'Линоленовая', 'Арахиновая', 'Бегеновая']
            predictions = {}
            for i, a in enumerate(acids):
                predictions[a] = (pred_ingr[i] + pred_nutr[i]) / 2.0
            self.current_predictions = predictions
            # Заполняем таблицу предсказаний
            self.pred_table.setRowCount(len(predictions))
            # Рассчитываем уровень по ГОСТу для каждого значения
            gost_levels = check_fatty_acid_ranges([predictions[a] for a in acids])
            for row, (acid, value) in enumerate(predictions.items()):
                self.pred_table.setItem(row, 0, QTableWidgetItem(FATTY_ACID_NAMES.get(acid, acid)))
                self.pred_table.setItem(row, 1, QTableWidgetItem(f"{value:.2f}"))
                self.pred_table.setItem(row, 2, QTableWidgetItem(gost_levels[row]))
                self.pred_table.setItem(row, 3, QTableWidgetItem(
                    "Высокая" if value > 0.1 else "Средняя"
                ))
            print(predictions)
            self.distribution_points = predictions
        except Exception as e:
            raise e

        """# Сохраняем предсказания в БД, если есть diet_id
        if self.current_diet_id:
            prediction_id = self.db.add_prediction(self.current_diet_id, predictions)
            self.statusBar().showMessage(f"Предсказания сгенерированы и сохранены (ID: {prediction_id})")
        else:
            # Если нет diet_id, создаем новый рацион
            diet_id = self.db.add_diet(
                name=f"Рацион из предсказаний {datetime.now().strftime('%Y-%m-%d %H:%M')}",
            )
            self.current_diet_id = diet_id
            prediction_id = self.db.add_prediction(diet_id, predictions)
            self.statusBar().showMessage(
                f"Предсказания сгенерированы и сохранены (Diet ID: {diet_id}, Prediction ID: {prediction_id})")
        # Также обновляем таблицу результатов
        self.update_results_table()"""

    def update_results_table(self):
        """Обновление таблицы результатов"""
        try:
            if not self.current_predictions:
                return
            self.results_pred_table.setRowCount(len(self.current_predictions))
            acids_order = ['Масляная', 'Капроновая', 'Каприловая', 'Каприновая', 'Деценовая', 'Лауриновая',
                           'Миристиновая', 'Миристолеиновая', 'Пальмитиновая', 'Пальмитолеиновая',
                           'Стеариновая', 'Олеиновая', 'Линолевая', 'Линоленовая', 'Арахиновая', 'Бегеновая']
            values_in_order = [self.current_predictions.get(a, 0.0) for a in acids_order]
            gost_levels = check_fatty_acid_ranges(values_in_order)
            for row, acid in enumerate(acids_order):
                value = self.current_predictions.get(acid, 0.0)
                self.results_pred_table.setItem(row, 0, QTableWidgetItem(FATTY_ACID_NAMES.get(acid, acid)))
                self.results_pred_table.setItem(row, 1, QTableWidgetItem(f"{value:.2f}"))
                # Форматируем «уровень по ГОСТу» и «уверенность»
                self.results_pred_table.setItem(row, 2, QTableWidgetItem(self._gost_level_text(gost_levels[row])))
                self.results_pred_table.setItem(row, 3, QTableWidgetItem(self._confidence_text(value)))
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка обновления: {str(e)}")


    def _gost_level_text(self, text: str) -> str:
        """Единый формат отображения уровня по ГОСТу в экспорте и печати."""
        try:
            if isinstance(text, str):
                if text.startswith('Ниже на '):
                    num = text.replace('Ниже на ', '')
                    return f"Ниже нормы на {num} %"
                if text.startswith('Выше на '):
                    num = text.replace('Выше на ', '')
                    return f"Выше нормы на {num} %"
            return text
        except Exception:
            return text

    def _confidence_text(self, value: float) -> str:
        """Единый формат отображения уверенности предсказания."""
        try:
            return "Высокая" if float(value) > 0.1 else "Средняя"
        except Exception:
            return "Средняя"


    def export_to_docx(self):
        """Экспорт результатов в DOCX"""
        try:
            if not self.current_predictions:
                QMessageBox.warning(self, "Предупреждение", "Сначала сгенерируйте предсказания!")
                return
            from docx import Document
            file_path, _ = QFileDialog.getSaveFileName(
                self, "Сохранить как", "результаты.docx", "Word files (*.docx)"
            )
            if file_path:
                doc = Document()
                doc.add_heading('Результаты предсказания жирнокислотного состава', 0)
                # 4 колонки: Кислота, Значение, Уровень по ГОСТу, Уверенность
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Light Grid Accent 1'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Жирная кислота'
                hdr_cells[1].text = 'Значение (%)'
                hdr_cells[2].text = 'Уровень по ГОСТу'
                hdr_cells[3].text = 'Уверенность'

                acids_order = ['Масляная', 'Капроновая', 'Каприловая', 'Каприновая', 'Деценовая', 'Лауриновая',
                               'Миристиновая', 'Миристолеиновая', 'Пальмитиновая', 'Пальмитолеиновая',
                               'Стеариновая', 'Олеиновая', 'Линолевая', 'Линоленовая', 'Арахиновая', 'Бегеновая']
                values_in_order = [self.current_predictions.get(a, 0.0) for a in acids_order]
                gost_levels = check_fatty_acid_ranges(values_in_order)

                for idx, acid in enumerate(acids_order):
                    value = values_in_order[idx]
                    row_cells = table.add_row().cells
                    row_cells[0].text = FATTY_ACID_NAMES.get(acid, acid)
                    row_cells[1].text = f"{value:.2f}"
                    row_cells[2].text = self._gost_level_text(gost_levels[idx])
                    row_cells[3].text = self._confidence_text(value)
                doc.save(file_path)
                QMessageBox.information(self, "Успех", f"Экспорт в DOCX завершен: {file_path}")
        except ImportError:
            QMessageBox.critical(self, "Ошибка", "Установите python-docx: pip install python-docx")
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка экспорта: {str(e)}")

    def export_to_pdf(self):
        """Экспорт результатов в PDF"""
        try:
            if not self.current_predictions:
                QMessageBox.warning(self, "Предупреждение", "Сначала сгенерируйте предсказания!")
                return

            from reportlab.lib.pagesizes import letter
            from reportlab.lib import colors
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont

            # Регистрируем шрифт с поддержкой кириллицы
            pdfmetrics.registerFont(TTFont('Arial', 'arial.ttf'))

            file_path, _ = QFileDialog.getSaveFileName(
                self, "Сохранить как", "результаты.pdf", "PDF files (*.pdf)"
            )
            if file_path:
                doc = SimpleDocTemplate(file_path, pagesize=letter)
                elements = []

                # Настраиваем стиль для русского текста
                styles = getSampleStyleSheet()
                title_style = ParagraphStyle(
                    'CustomTitle',
                    parent=styles['Title'],
                    fontName='Arial',
                    fontSize=16,
                    leading=20
                )
                normal_style = ParagraphStyle(
                    'CustomNormal',
                    parent=styles['Normal'],
                    fontName='Arial',
                    fontSize=12,
                    leading=14
                )

                elements.append(Paragraph("Результаты предсказания жирнокислотного состава", title_style))

                data = [['Жирная кислота', 'Значение (%)', 'Уровень по ГОСТу', 'Уверенность']]

                acids_order = ['Масляная', 'Капроновая', 'Каприловая', 'Каприновая', 'Деценовая', 'Лауриновая',
                               'Миристиновая', 'Миристолеиновая', 'Пальмитиновая', 'Пальмитолеиновая',
                               'Стеариновая', 'Олеиновая', 'Линолевая', 'Линоленовая', 'Арахиновая', 'Бегеновая']
                values_in_order = [self.current_predictions.get(a, 0.0) for a in acids_order]
                gost_levels = check_fatty_acid_ranges(values_in_order)

                for idx, acid in enumerate(acids_order):
                    value = values_in_order[idx]
                    data.append([
                        FATTY_ACID_NAMES.get(acid, acid),
                        f"{value:.2f}",
                        self._gost_level_text(gost_levels[idx]),
                        self._confidence_text(value)
                    ])

                table = Table(data)
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Arial'),
                    ('FONTSIZE', (0, 0), (-1, 0), 14),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('FONTNAME', (0, 1), (-1, -1), 'Arial'),  # Применяем шрифт ко всей таблице
                    ('FONTSIZE', (0, 1), (-1, -1), 10),
                ]))

                elements.append(table)
                doc.build(elements)
                QMessageBox.information(self, "Успех", f"Экспорт в PDF завершен: {file_path}")

        except ImportError:
            QMessageBox.critical(self, "Ошибка", "Установите reportlab и DejaVuSans: pip install reportlab")
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка экспорта: {str(e)}")

    def print_results(self):
        """Печать результатов"""
        try:
            if not self.current_predictions:
                QMessageBox.warning(self, "Предупреждение", "Сначала сгенерируйте предсказания!")
                return
            from PyQt6.QtPrintSupport import QPrinter, QPrintDialog
            from PyQt6.QtGui import QPainter
            printer = QPrinter()
            dialog = QPrintDialog(printer, self)
            if dialog.exec() == QPrintDialog.DialogCode.Accepted:
                painter = QPainter(printer)
                # Простая печать таблицы
                # Обновляем таблицу перед печатью, чтобы гарантировать актуальные значения и формат
                self.update_results_table()
                self.results_pred_table.render(painter)
                painter.end()
                QMessageBox.information(self, "Успех", "Печать завершена!")
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка печати: {str(e)}")

    def show_hist(self):
        """Показать распределение"""
        try:
            if not hasattr(self, 'distribution_points') or not self.distribution_points:
                self.canvas.axes.clear()
                self.canvas.axes.text(0.5, 0.5, 'Нет данных', ha='center', va='center')
                self.canvas.draw()
                return

            # Получаем ключи и значения из словаря
            labels = list(self.distribution_points.keys())
            values = list(self.distribution_points.values())

            # Проверяем, что списки не пусты
            if len(labels) == 0 or len(values) == 0:
                self.canvas.axes.clear()
                self.canvas.axes.text(0.5, 0.5, 'Нет данных', ha='center', va='center')
                self.canvas.draw()
                return

            self.canvas.axes.clear()
            bars = self.canvas.axes.bar(labels, values, alpha=0.7)

            # Подписываем значения сверху столбцов
            for bar, value in zip(bars, values):
                height = bar.get_height()
                self.canvas.axes.text(
                    bar.get_x() + bar.get_width() / 2., height,
                    f'{value:.2f}',
                    ha='center', va='bottom'
                )

            self.canvas.axes.set_xlabel('Жирные кислоты')
            self.canvas.axes.set_ylabel('Значение (%)')
            self.canvas.axes.set_title('Распределение жирных кислот')
            self.canvas.axes.tick_params(axis='x', rotation=45)
            self.canvas.figure.tight_layout()
            self.canvas.draw()

        except Exception as e:
            from PyQt6.QtWidgets import QMessageBox
            QMessageBox.critical(self, "Ошибка", f"Ошибка графика: {str(e)}")

    def show_distribution(self):
        """Показать распределение"""
        try:
            diets = self.db.get_all_diets()
            all_analyses = []
            for diet in diets:
                analyses = self.db.get_analysis_for_diet(diet['id'])
                all_analyses.extend(analyses)
            if not all_analyses:
                self.canvas.axes.clear()
                self.canvas.axes.text(0.5, 0.5, 'Нет данных', ha='center', va='center')
                self.canvas.draw()
                return
            acids = ['lauric_acid', 'palmitic_acid', 'stearic_acid',
                     'oleic_acid', 'linoleic_acid', 'linolenic_acid']
            names = ['Лауриновая', 'Пальмитиновая', 'Стеариновая',
                     'Олеиновая', 'Линолевая', 'Линоленовая']
            data_means = []
            labels = []
            for acid, name in zip(acids, names):
                values = [a[acid] for a in all_analyses if a[acid]]
                if values:
                    data_means.append(np.mean(values))
                    labels.append(name)
            self.canvas.axes.clear()
            self.canvas.axes.bar(labels, data_means, alpha=0.7, color='steelblue')
            self.canvas.axes.set_xlabel('Жирные кислоты')
            self.canvas.axes.set_ylabel('Среднее значение (%)')
            self.canvas.axes.set_title('Распределение жирных кислот')
            self.canvas.axes.tick_params(axis='x', rotation=45)
            self.canvas.figure.tight_layout()
            self.canvas.draw()
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка графика: {str(e)}")

    def show_trends(self):
        """Показать тренды"""
        self.canvas.axes.clear()
        self.canvas.axes.text(0.5, 0.5, 'Функция трендов в разработке',
                              ha='center', va='center', fontsize=14)
        self.canvas.draw()

    # -------------------- Галерея заготовленных графиков --------------------
    def on_tab_changed(self, index: int):
        """Автоматически обновлять миниатюры при входе во вкладку результатов."""
        try:
            if hasattr(self, 'results_tab_index') and index == self.results_tab_index:
                self.load_prepared_images()
        except Exception:
            pass

    def clear_prepared_grid(self):
        """Очистка сетки миниатюр."""
        try:
            while self.prepared_grid.count():
                item = self.prepared_grid.takeAt(0)
                widget = item.widget()
                if widget is not None:
                    widget.setParent(None)
                    widget.deleteLater()
        except Exception:
            pass

    def load_prepared_images(self):
        """Загрузка изображений из папки visuals и отображение миниатюр."""
        try:
            # Если секция ещё не инициализирована (например, вкладка не создана)
            if not hasattr(self, 'prepared_grid'):
                return
            self.clear_prepared_grid()
            if not os.path.isdir(self.visuals_dir):
                placeholder = QLabel("Папка с графиками не найдена: " + self.visuals_dir)
                placeholder.setAlignment(Qt.AlignmentFlag.AlignCenter)
                self.prepared_grid.addWidget(placeholder, 0, 0)
                return
            entries = sorted(os.listdir(self.visuals_dir))
            image_exts = {'.png', '.jpg', '.jpeg', '.webp', '.bmp', '.gif'}
            image_paths = [
                os.path.join(self.visuals_dir, name)
                for name in entries
                if os.path.splitext(name)[1].lower() in image_exts
            ]
            if not image_paths:
                placeholder = QLabel("Нет изображений в папке visuals")
                placeholder.setAlignment(Qt.AlignmentFlag.AlignCenter)
                self.prepared_grid.addWidget(placeholder, 0, 0)
                return
            columns = 3
            row = 0
            col = 0
            thumb_max_w, thumb_max_h = 320, 220
            for path in image_paths:
                # Контейнер одного элемента
                item_widget = QWidget()
                vbox = QVBoxLayout(item_widget)
                vbox.setContentsMargins(6, 6, 6, 6)
                vbox.setSpacing(6)
                pixmap = QPixmap(path)
                if not pixmap.isNull():
                    scaled = pixmap.scaled(
                        thumb_max_w, thumb_max_h,
                        Qt.AspectRatioMode.KeepAspectRatio,
                        Qt.TransformationMode.SmoothTransformation
                    )
                    img_label = ClickableLabel(path)
                    img_label.setPixmap(scaled)
                    img_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
                    img_label.clicked.connect(self.open_image_viewer)
                    vbox.addWidget(img_label)
                else:
                    broken = QLabel("Не удалось загрузить изображение")
                    broken.setAlignment(Qt.AlignmentFlag.AlignCenter)
                    vbox.addWidget(broken)
                name_label = QLabel(os.path.basename(path))
                name_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
                name_label.setWordWrap(True)
                vbox.addWidget(name_label)
                self.prepared_grid.addWidget(item_widget, row, col)
                col += 1
                if col >= columns:
                    col = 0
                    row += 1
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка загрузки изображений: {str(e)}")

    def open_image_viewer(self, image_path: str):
        """Открыть диалог с полноразмерным просмотром изображения."""
        try:
            dialog = QDialog(self)
            dialog.setWindowTitle(os.path.basename(image_path))
            dialog.resize(1000, 700)
            layout = QVBoxLayout(dialog)
            scroll = QScrollArea()
            scroll.setWidgetResizable(True)
            container = QWidget()
            v = QVBoxLayout(container)
            v.setContentsMargins(0, 0, 0, 0)
            v.setSpacing(0)
            lbl = QLabel()
            lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
            pm = QPixmap(image_path)
            if pm.isNull():
                lbl.setText("Не удалось загрузить изображение")
            else:
                # Подгоняем под текущее окно при открытии
                area_w = dialog.width() - 40
                area_h = dialog.height() - 60
                scaled = pm.scaled(
                    max(100, area_w), max(100, area_h),
                    Qt.AspectRatioMode.KeepAspectRatio,
                    Qt.TransformationMode.SmoothTransformation
                )
                lbl.setPixmap(scaled)
            v.addWidget(lbl)
            scroll.setWidget(container)
            layout.addWidget(scroll)
            btn_box = QHBoxLayout()
            close_btn = QPushButton("Закрыть")
            close_btn.clicked.connect(dialog.accept)
            btn_box.addStretch()
            btn_box.addWidget(close_btn)
            layout.addLayout(btn_box)
            dialog.exec()
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка просмотра изображения: {str(e)}")


def main():
    app = QApplication(sys.argv)
    app.setStyle('Fusion')
    window = MainWindow()
    window.show()
    return app.exec()


if __name__ == "__main__":
    main()
